from concurrent.futures import ThreadPoolExecutor, as_completed
import streamlit as st
import sys 
import os
import sqlite3
import json
from langchain_tavily import TavilySearch, TavilyExtract
from langchain_openai import ChatOpenAI
from langgraph.graph import StateGraph
from langchain_core.output_parsers import StrOutputParser
from langchain.prompts import PromptTemplate
from langgraph.graph import END
from pydantic import BaseModel
from typing import List, Dict
from codecarbon import OfflineEmissionsTracker
from dotenv import load_dotenv
from docx import Document
from io import BytesIO

from datetime import datetime

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from scripts.get_carbon_data import get_codecarbon_estimate

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

load_dotenv()

st.set_page_config(layout="wide")

st.markdown("""
    <style>
        body {
            background-color: #F7FFF7;
        }
        .block-container {
            padding-top: 2rem;
        }
        .info-card {
            background-color: white;
            border-radius: 12px;
            padding: 20px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.05);
            text-align: left;
        }
        .info-card h4 {
            margin-bottom: 10px;
        }
        .info-card p {
            margin-bottom: 20px;
        }

        /* Style Streamlit buttons */
        div.stButton > button:first-child {
            background-color: #66ccff;
            color: black;
            border: none;
            border-radius: 8px;
            padding: 8px 20px;
            font-size: 16px;
            cursor: pointer;
        }
        div.stButton > button:first-child:hover {
            background-color: #4dc3ff;
        }

        .small-card {
            background-color: #f9f9f9;
            border-radius: 10px;
            padding: 20px;
            height: 160px;
        }
        .small-card h5 {
            margin: 0 0 10px 0;
        }

        /* Custom styling for summary card */
        .summary-card {
            background-color: #f0f8ff;
            border-radius: 10px;
            padding: 20px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            text-align: left;
            margin-bottom: 20px;
        }

        /* Custom styling for references card */
        .references-card {
            background-color: #f0f8ff;
            border-radius: 10px;
            padding: 20px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            text-align: left;
            margin-bottom: 20px;
        }

        /* Back button styling */
        .stButton>button.back {
            background-color: #66ccff;
            color: black;
            border: none;
            border-radius: 6px;
            padding: 6px 15px;
            font-size: 16px;
            margin-top: 100px;
        }
    </style>
""", unsafe_allow_html=True)


class ResearchState(BaseModel):
    query: str
    search_results: List[Dict] = []
    extracted_docs: List[Dict] = []
    summaries: List[Dict] = []
    sources: List[Dict] = []

# SQLite Notebook
db = sqlite3.connect('notebook.db')
c = db.cursor()
c.execute('''
    CREATE TABLE IF NOT EXISTS notebook (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        question TEXT,
        summary TEXT,
        title TEXT,
        url TEXT,
        timestamp TEXT,
        UNIQUE(question, title, url)
    )
''')

try:
    c.execute('ALTER TABLE notebook ADD COLUMN notes TEXT')
    db.commit()
except sqlite3.OperationalError:
    # The column already exists
    pass
    
db.commit()

llm_model = "gpt-4.1-nano"

# Back button
top_col1, _, _ = st.columns([0.1, 0.8, 0.1])
with top_col1:
    st.markdown("<div style='margin-top: 30px;'></div>", unsafe_allow_html=True)
    if st.button("Back", key="back"):
        st.session_state.pop("GHG_selected_question", None)
        st.switch_page("research_topic_page.py")


st.markdown('<h1 style="text-align: center;"> AI Environmental Impact Research </h1>', unsafe_allow_html=True)

research_topic_col, codecarbon_col = st.columns([3, 1])

with research_topic_col:
    st.markdown("**Research Topic**")

    if "GHG_user_question" in st.session_state:
        st.markdown("> **" + st.session_state["GHG_user_question"] + "**")
    elif "GHG_selected_question" in st.session_state:
        st.markdown("> **" + st.session_state["GHG_selected_question"] + "**")
        
    st.markdown("**Research Domains**")
    included_domains = ["ieeexplore.ieee.org", "springeropen.com", "scienceopen.com", "nature.com", "sciencedirect.com"]    # Peer-reviewed domains for Tavily search. "arxiv.org" not peer-reviewed.
    st.markdown("\n".join(f"- {domain}" for domain in included_domains))

with codecarbon_col: # from Codecarbon emissions in kg Co2eq , energy in kWh, this will be converted
    codecarbon_data = get_codecarbon_estimate()
    
    emissions = codecarbon_data["emissions"]
    energy_consumed = codecarbon_data["energy_consumed"]
    timestamp = codecarbon_data["timestamp"]
    
    dt = datetime.strptime(timestamp, "%Y-%m-%dT%H:%M:%S")
    data_retrieval_time = dt.strftime("%B %d, %Y at %I:%M %p")
    
    st.markdown("<br>", unsafe_allow_html=True)
    with st.expander("**What is the environmental impact of this service?**"):
        st.markdown(f"Sources are retrieved with Tavily Search and Tavily Extract, using only the trusted domains stated under 'Research Domains'. This service also uses ChatOpenAI from Langchain with the model {llm_model} to generate research topics and generate a summarization of each article.")
        st.markdown(f"To estimate the energy consumed (sum of cpu_energy, gpu_energy and ram_energy) and carbon emissions on a local machine, the summarization step with the help of {llm_model} was tracked with CodeCarbon, and the values are displayed below:")
        st.markdown(f'**Energy use of LLM summarization of this service: {energy_consumed:.4f} Wh.**')
        st.markdown(f'**Carbon emissions of LLM summarization of this service: {emissions:.4f} g CO2eq.**')
        st.markdown(f'These values were calculated on the {data_retrieval_time} Swedish time, on an Apple M1 Pro and is just an estimation.')
            

# SUMMARY SECTION

@st.cache_data(show_spinner="Searching Tavily...")
def cached_search(query):
    result = search_tool.invoke({"query": query})
    return result["results"]


@st.cache_data(show_spinner="Extracting content...")
def cached_extract(url):
    return extract_tool.invoke({"urls": [url]})


@st.cache_data(show_spinner="Summarizing content...")
def cached_summarize(prompt_input):
    return llm.invoke(prompt_input)

llm = ChatOpenAI (
# model="gpt-4o",
# model="gpt-4o-mini",
model=llm_model,
# temperature=0.7,       
temperature = 0       # reduce randomness and wasted tokens
)

# Tavily tools
search_tool = TavilySearch(
    max_results=6,
    include_domains=included_domains  
)
extract_tool = TavilyExtract(include_images=False)

graph = StateGraph(ResearchState)
    

def search_node(state):
    query = state.query
    search_results = cached_search(query)
    
    return {"search_results": search_results}

graph.add_node("search", search_node)

def extract_node(state: ResearchState):
    url_to_title = {result['url']: result.get('title', '') for result in state.search_results}
    url_to_content = {result['url']: result.get('content', '') for result in state.search_results}
    extracted = []
    for url in url_to_title.keys():
        extracted_doc = cached_extract(url)
        
        # TavilyExtract returns a list under "results"
        if isinstance(extracted_doc, dict) and "results" in extracted_doc:
            raw_content = extracted_doc["results"][0]["raw_content"]       # shorten content?
        else:
            raw_content = str(extracted_doc)
            
        if not raw_content:  # If extraction is empty, get short content from TavilySearch
            raw_content = url_to_content.get(url, "")
            
        extracted.append({
            "url": url,
            "title": url_to_title.get(url, "Untitled"),  
            "text": raw_content  
        })
    return {"extracted_docs": extracted}

graph.add_node("extract", extract_node)

def summarize_article(source, query):
    # Uncomment for tracking emissions:
    # tracker = OfflineEmissionsTracker(country_iso_code="SWE")
    # tracker.start()
    summary_prompt = PromptTemplate.from_template(
        """ Research topic: {query}
            Name of article: {article_title}
            Extracted text from research article: {research_text}
            
            You are a helpful research assistant that aims to summarize research papers to someone who wants to understand the main point of the article, based on the given research topic. 
            Write a summary according to the instructions below. 
            - Write a short summary of the article, about 2-3 sentences.
            - If quantified data is included in the research paper, include this data in the summary.
            - Be specific in the summary rather than vague.
            - Do NOT make up any new information, make sure that you stick to information in the extracted text.
    """
    )
    
    create_summary_chain = summary_prompt | llm | StrOutputParser()
    tokens = create_summary_chain.invoke({
        "query": query,
        "article_title": source["title"],
        "research_text": source["text"]     # truncate [:4000]?  
    })
    
    # tracker.stop()
    
    return {
        "title": source["title"],
        "url": source["url"],
        "summary": tokens
    }
    
def summarize_node(state):
    summaries = []
    with ThreadPoolExecutor(max_workers=2) as executor:
        futures = [
            executor.submit(summarize_article, source, state.query)
            for source in state.extracted_docs
        ]
        progress_placeholder = st.empty()
        progress = progress_placeholder.progress(0)
        total = len(futures)
        done = 0

        for future in as_completed(futures):
            result = future.result()
            summaries.append(result)
            done += 1
            progress.progress(done / total)
            
        progress_placeholder.empty()  # hide the progress bar when the progress bar is finished

    sources = []
    for i, s in enumerate(summaries, start=1):
        sources.append({
            "index": i,
            "url": s["url"],
            "title": s["title"]
        })

    return {
        "summaries": summaries,
        "sources": sources
    }
    
    
graph.add_node("summarize", summarize_node)
    
# Edges
graph.add_edge("search", "extract")
graph.add_edge("extract", "summarize")
graph.add_edge("summarize", END)

graph.set_entry_point("search")

# Graph: (search) → (extract) → (summarize) → END
app = graph.compile()

selected_question = st.session_state.get("GHG_user_question") or st.session_state.get("GHG_selected_question")

inputs = {"query": selected_question}

if "GHG_result" not in st.session_state or st.session_state["GHG_result"]["query"] != selected_question:
    result = app.invoke(inputs)  
    st.session_state["GHG_result"] = result  
    
else:
    result = st.session_state["GHG_result"]

final_state = ResearchState(
    query = result["query"],
    search_results = result["search_results"],
    extracted_docs = result["extracted_docs"],
    summaries = result["summaries"],
    sources = result["sources"]
)

title_col, select_all_col = st.columns([0.90, 0.1])

with title_col:
    st.markdown('<h2 style="text-align: left;"> Research Articles </h2>', unsafe_allow_html=True)
    st.markdown(':gray[To save research articles to the Research Notebook, click one or several of the checkboxes below]')
    
with select_all_col:
    if st.button("Select all"):
        for i in range(len(final_state.summaries)):
            st.session_state[f"checkbox_{i}"] = True
    
    
c.execute('SELECT url FROM notebook')
saved_urls = {row[0] for row in c.fetchall()}
 
selected_articles = []
for i, source in enumerate(final_state.summaries):
    key = f"checkbox_{i}"
    if key not in st.session_state:
        st.session_state[key] = source["url"] in saved_urls
        
    col1, col2 = st.columns([0.04, 0.96])  

    with col1:
        checked = st.checkbox("🔗", key=key)     # value=True if every box should be preselected to true

    with col2:
        st.markdown(
            f"<div style='margin-top: 7px'><a href='{source['url']}' target='_blank'> {source['title']}</a></div>",
            unsafe_allow_html=True)

    st.markdown(source["summary"])
    st.markdown("---")
    
    # if checked and not is_saved:
    if st.session_state[key] and source["url"] not in saved_urls:
        c.execute('''
            INSERT OR IGNORE INTO notebook (question, summary, title, url, timestamp)
            VALUES (?, ?, ?, ?, ?)
        ''', (final_state.query, source["summary"], source["title"], source["url"], datetime.now().isoformat()))
        
        db.commit()
    
    # elif not checked and is_saved:
    elif not st.session_state[key] and source["url"] in saved_urls:
        c.execute('DELETE FROM notebook WHERE url = ?', (source["url"],))
        db.commit()
        # st.info(f"Removed: {source['title']}")
        st.info("Removed article from Research Notebook")
    
    
with st.sidebar:
    if st.button("Clear Research Notebook"):
        c.execute('DELETE FROM notebook')
        db.commit()
        st.success("Your Research Notebook has been cleared!")
        
    # View notebook 
    st.markdown('<h2 style="text-align: left;"> Research Notebook </h2>', unsafe_allow_html=True)

    # Make it possible to search for key terms in summary, title or url
    search_term = st.text_input("Search Notebook")  
    if search_term:
        c.execute('''
            SELECT question, summary, title, url, timestamp 
            FROM notebook 
            WHERE question LIKE ? OR summary LIKE ? OR title LIKE ?
            ORDER BY timestamp DESC
        ''', (f'%{search_term}%', f'%{search_term}%',  f'%{search_term}%'))
       
    else:
        c.execute('SELECT question, summary, title, url, timestamp FROM notebook ORDER BY timestamp DESC')
        
        
    rows = c.fetchall()
    if rows:
        for i, row in enumerate(rows):
            st.markdown(f"**{row[0]}**  \n_Saved: {row[4]}_", unsafe_allow_html=True)
            st.markdown(f"[🔗 {row[2]}]({row[3]})", unsafe_allow_html=True)
            st.markdown(row[1])
            
            # Load existing note from DB
            c.execute('SELECT notes FROM notebook WHERE url = ?', (row[3],))
            note_row = c.fetchone()
            existing_note = note_row[0] if note_row and note_row[0] else ""

            # Editable text area
            note_input = st.text_area(
                label="Notes",
                value=existing_note,
                key=f"note_sidebar_{i}",
                height=100,
                label_visibility="visible",
                placeholder="Why did you save this article?"
            )

            # Save back to DB if changed
            if note_input != existing_note:
                c.execute('UPDATE notebook SET notes = ? WHERE url = ?', (note_input, row[3]))
                db.commit()
                st.success(f"Note saved for: {row[2]}")

            st.markdown("---")
    else:
        if search_term:
            # If research is saved, but the search query has no match
            st.info("No search results matched your query. Try something else!")
            
        else:  
            st.info("No entries yet. Save some research!")


    c.execute('SELECT question, summary, title, url, timestamp, notes FROM notebook')
    rows = c.fetchall()
    export = [
        {"question": r[0], "summary": r[1], "title": r[2], "url": r[3], "timestamp": r[4], "notes": r[5]}
        for r in rows
    ]
    
    
    def create_word_doc(json_data, title="Research Notebook"):
        doc = Document()
        doc.add_heading(title, level=1)
        for i, item in enumerate(json_data, 1):
            p = doc.add_paragraph()
            p.add_run("\n")
            p.add_run(f"{i}. Title: ").bold = True
            p.add_run(item['title'])
            p = doc.add_paragraph()
            p.add_run("Topic: ").bold = True
            p.add_run(item['question'])
            p = doc.add_paragraph()
            p.add_run("Summary: ").bold = True
            p.add_run(item['summary'])
            p = doc.add_paragraph()
            p.add_run("URL: ").bold = True
            p.add_run({item['url']})
            
            try:
                dt = datetime.fromisoformat(item['timestamp'])
                formatted_time = dt.strftime("%B %d, %Y at %I:%M %p")
            except Exception:
                formatted_time = item['timestamp']  # fallback 

            p = doc.add_paragraph()
            p.add_run("Timestamp: ").bold = True
            p.add_run(formatted_time)
            
            p = doc.add_paragraph()
            p.add_run("Notes: ").bold = True
            p.add_run(item['notes'])
            p.add_run("\n\n")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    word_file = create_word_doc(export)

    st.download_button(
        "⬇ Download Notebook File",
        word_file,
        file_name="notebook.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
            
